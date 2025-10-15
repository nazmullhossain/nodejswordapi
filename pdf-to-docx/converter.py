import sys
import os
import tempfile
import traceback

def convert_pdf_to_docx(pdf_path, docx_path):
    """
    Convert PDF to DOCX with fallback options for Railway hosting
    """
    try:
        print(f"[INFO] Starting conversion: {pdf_path}")
        
        if not os.path.exists(pdf_path):
            print(f"[ERROR] Input file not found: {pdf_path}")
            return False
        
        # Use temp directory for processing (Railway compatible)
        temp_dir = tempfile.gettempdir()
        temp_output = os.path.join(temp_dir, os.path.basename(docx_path))
        
        print(f"[INFO] Input: {pdf_path}")
        print(f"[INFO] Output: {docx_path}")
        print(f"[INFO] Temp: {temp_output}")
        
        success = False
        
        # Method 1: Try pdf2docx with parse (most reliable)
        try:
            from pdf2docx import parse
            print("[INFO] Attempting conversion with pdf2docx...")
            parse(pdf_path, temp_output)
            success = True
            print("[SUCCESS] pdf2docx conversion completed")
        except Exception as e:
            print(f"[WARNING] pdf2docx failed: {str(e)}")
        
        # Method 2: Fallback to basic text extraction
        if not success:
            try:
                print("[INFO] Falling back to basic text extraction...")
                import fitz
                from docx import Document
                
                pdf_document = fitz.open(pdf_path)
                doc = Document()
                total_pages = len(pdf_document)
                
                print(f"[INFO] Processing {total_pages} pages...")
                
                for page_num in range(total_pages):
                    page = pdf_document.load_page(page_num)
                    text = page.get_text()
                    
                    if text.strip():
                        # Add page header
                        header = doc.add_paragraph()
                        header.add_run(f"Page {page_num + 1}").bold = True
                        
                        # Add content with basic formatting
                        for line in text.split('\n'):
                            if line.strip():
                                doc.add_paragraph(line.strip())
                    
                    if (page_num + 1) % 10 == 0 or (page_num + 1) == total_pages:
                        progress = int(((page_num + 1) / total_pages) * 100)
                        print(f"[PROGRESS] Page {page_num + 1}/{total_pages} ({progress}%)")
                
                doc.save(temp_output)
                pdf_document.close()
                success = True
                print("[SUCCESS] Basic conversion completed")
                
            except Exception as e:
                print(f"[ERROR] Basic conversion also failed: {str(e)}")
                return False
        
        # Move file to final location if temp was used
        if success and temp_output != docx_path:
            import shutil
            os.makedirs(os.path.dirname(docx_path), exist_ok=True)
            shutil.move(temp_output, docx_path)
            print(f"[INFO] File moved to final location: {docx_path}")
        
        # Final verification
        if os.path.exists(docx_path):
            file_size = os.path.getsize(docx_path)
            print(f"[SUCCESS] Final file: {docx_path} ({file_size} bytes)")
            return True
        else:
            print("[ERROR] Final output file not found")
            return False
            
    except Exception as e:
        print(f"[FATAL ERROR] Conversion failed: {str(e)}")
        traceback.print_exc()
        return False

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python converter.py <input_pdf> <output_docx>")
        sys.exit(1)
        
    pdf_path = sys.argv[1]
    docx_path = sys.argv[2]
    
    success = convert_pdf_to_docx(pdf_path, docx_path)
    
    if success:
        print("CONVERSION_STATUS:SUCCESS")
        print(f"OUTPUT_FILE:{docx_path}")
    else:
        print("CONVERSION_STATUS:FAILED")
    
    sys.exit(0 if success else 1)